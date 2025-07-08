import streamlit as st
import pandas as pd
import plotly.express as px
from datetime import datetime
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

def show_can_dashboard():
    st.title("🛠️ CAN - Component Alert Notice Dashboard")

    # قراءة البيانات
    df_raw = pd.read_excel("CAN.xlsx")
    df = df_raw.loc[:, ~df_raw.columns.str.contains("^Unnamed")]
    df.columns = df.columns.str.strip()

    # إعداد الفلاتر
    years = sorted(df["YEAR"].dropna().unique())
    quarters = sorted(df["QUARTER NO"].dropna().unique())
    ac_types = sorted(df["A/C TYPE"].dropna().unique())

    left_col, right_col = st.columns([1, 2])
    with left_col:
        st.subheader("🔍 Filters")
        selected_years = st.multiselect("Select Year(s)", years, default=years)
        selected_quarters = st.multiselect("Select Quarter(s)", quarters, default=quarters)
        selected_types = st.multiselect("Select A/C Type(s)", ac_types, default=ac_types)

    filtered_df = df[
        df["YEAR"].isin(selected_years) &
        df["QUARTER NO"].isin(selected_quarters) &
        df["A/C TYPE"].isin(selected_types)
    ]

    with right_col:
        st.markdown("### Filtered CAN Data")
        # 🔍 مربع بحث للجدول
        search_term = st.text_input("Search inside table (by any keyword)", "")
        filtered_display_df = filtered_df.copy()

        if search_term:
            filtered_display_df = filtered_display_df[
                filtered_display_df.apply(lambda row: row.astype(str).str.contains(search_term, case=False).any(), axis=1)
            ]

        st.dataframe(filtered_display_df, use_container_width=True)

        # 💡 إجمالي عدد الـ CAN
        st.markdown(f"✅ **Total CAN Entries (after filter):** `{len(filtered_display_df)}`")

        # تحميل Excel
        excel_buffer = BytesIO()
        with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
            filtered_display_df.to_excel(writer, sheet_name="Filtered CAN", index=False)
        excel_buffer.seek(0)
        st.download_button(
            label="📥 Download Filtered Data as Excel",
            data=excel_buffer,
            file_name="Filtered_CAN_Data.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    # ✅ جدول عدد CAN لكل P/N + بحث داخلي + تفاصيل
    st.markdown("### 📦 Count of CAN per Part Number")

    part_counts = filtered_df["P/N"].value_counts().reset_index()
    part_counts.columns = ["Part Number", "No. of CAN"]

    part_search = st.text_input("🔍 Search by Part Number", "")
    if part_search:
        part_counts = part_counts[part_counts["Part Number"].str.contains(part_search, case=False, na=False)]

    st.dataframe(part_counts, use_container_width=True)

    # 🧩 اختيار Part Number لعرض التفاصيل
    if not part_counts.empty:
        selected_part = st.selectbox("Select a Part Number to view CAN details", ["-- Select --"] + part_counts["Part Number"].tolist())

        if selected_part and selected_part != "-- Select --":
            st.markdown(f"### 📄 CAN Details for P/N: `{selected_part}`")
            part_details = filtered_df[filtered_df["P/N"] == selected_part]
            st.dataframe(part_details, use_container_width=True)



   

    # 🚨 جدول: P/N فيها تجاوز للـ Alert
    st.markdown("### 🚨 Parts Exceeding REMOVAL ALERT")
    if "REMOVAL RATE" in filtered_df.columns and "REMOVAL ALERT" in filtered_df.columns:
        exceed_df = filtered_df[filtered_df["REMOVAL RATE"] > filtered_df["REMOVAL ALERT"]]
        exceed_grouped = exceed_df.groupby("P/N").size().reset_index(name="Exceed Count")
        exceed_grouped = exceed_grouped.sort_values(by="Exceed Count", ascending=False)
        st.dataframe(exceed_grouped, use_container_width=True)
    else:
        st.warning("Columns 'REMOVAL RATE' or 'REMOVAL ALERT' not found.")

    st.markdown("---")
    st.subheader("📊 Visual Analytics")

    # ===== Chart 1 =====
    st.markdown("#### 🔧 Number of Removals by ATA")
    sort_removal = st.selectbox("Sort Order (Removals)", ["Descending", "Ascending"], index=0)
    top_removal = st.selectbox("Show Top (Removals)", ["All", "Top 3", "Top 6", "Top 10"], index=0)

    removal_by_ata = filtered_df.groupby("ATA")["NO OF REMOVAL"].sum().reset_index()
    removal_by_ata = removal_by_ata.sort_values("NO OF REMOVAL", ascending=(sort_removal == "Ascending"))
    if top_removal != "All":
        n = int(top_removal.split()[1])
        removal_by_ata = removal_by_ata.head(n)

    fig1 = px.bar(
        removal_by_ata,
        x="ATA",
        y="NO OF REMOVAL",
        text="NO OF REMOVAL",
        color_discrete_sequence=["#1f77b4"],
        category_orders={"ATA": removal_by_ata["ATA"].tolist()}
    )
    fig1.update_traces(textposition="outside")
    fig1.update_layout(
        yaxis_tickformat=",d",
        xaxis_tickmode='linear',
        xaxis_tickangle=-45
    )
    st.plotly_chart(fig1, use_container_width=True)

    

    # ===== Chart 3 =====
    st.markdown("#### 📈 CAN Distribution by ATA (Pie Chart)")
    top_pie = st.selectbox("Show Top (Pie)", ["All", "Top 3", "Top 6", "Top 10"], index=0)

    removal_by_ata_pie = filtered_df.groupby("ATA")["NO OF REMOVAL"].sum().reset_index()
    removal_by_ata_pie = removal_by_ata_pie.sort_values("NO OF REMOVAL", ascending=False)
    if top_pie != "All":
        n = int(top_pie.split()[1])
        removal_by_ata_pie = removal_by_ata_pie.head(n)

    if not removal_by_ata_pie.empty:
        fig3 = px.pie(
            removal_by_ata_pie,
            names="ATA",
            values="NO OF REMOVAL",
            color_discrete_sequence=px.colors.sequential.Blues
        )
        fig3.update_traces(textinfo='label+percent+value')
        st.plotly_chart(fig3, use_container_width=True)
    else:
        st.warning("No data available to display the Pie Chart.")

    # ===== Chart 4 =====
    st.markdown("#### 📌 Number of CAN Entries per ATA Chapter")
    sort_can = st.selectbox("Sort Order (CAN Count)", ["Descending", "Ascending"], index=0)
    top_can = st.selectbox("Show Top (CAN Count)", ["All", "Top 3", "Top 6", "Top 10"], index=0)

    can_count_by_ata = filtered_df.groupby("ATA").size().reset_index(name="CAN Count")
    can_count_by_ata = can_count_by_ata.sort_values("CAN Count", ascending=(sort_can == "Ascending"))
    if top_can != "All":
        n = int(top_can.split()[1])
        can_count_by_ata = can_count_by_ata.head(n)

    if not can_count_by_ata.empty:
        fig4 = px.bar(
            can_count_by_ata,
            x="CAN Count",
            y="ATA",
            orientation="h",
            text="CAN Count",
            color_discrete_sequence=["#2ca02c"],
            category_orders={"ATA": can_count_by_ata["ATA"].tolist()}
        )
        fig4.update_traces(textposition="outside")
        fig4.update_layout(
            xaxis_tickformat=",d",
            yaxis_tickmode='linear',
            yaxis_tickfont=dict(size=12)
        )
        st.plotly_chart(fig4, use_container_width=True)
    else:
        st.warning("No data available to display CAN count chart.")

    # ===== Export Reports =====
    st.markdown("---")
    st.subheader("📥 Export Report")

    # Word Export
    if st.button("📄 Download Word Report"):
        doc = Document()
        doc.add_heading("EGYPTAIR M&E", level=1)
        doc.add_paragraph("Technical Services Directorate")
        doc.add_paragraph("Reliability Department")
        doc.add_paragraph(f"Date: {datetime.today().strftime('%Y-%m-%d')}")
        doc.add_paragraph("Component Alert Notice (CAN) Report")

        doc.add_paragraph(f"Filters Used:")
        doc.add_paragraph(f"  • Years: {', '.join(map(str, selected_years))}")
        doc.add_paragraph(f"  • Quarters: {', '.join(map(str, selected_quarters))}")
        doc.add_paragraph(f"  • A/C Types: {', '.join(selected_types)}")

        doc.add_paragraph("\nSummary Data:")
        doc.add_paragraph(f"  • Total Records: {len(filtered_df)}")
        doc.add_paragraph(f"  • Unique ATA Chapters: {filtered_df['ATA'].nunique()}")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="📥 Click to download report (.docx)",
            data=buffer,
            file_name="EGYPTAIR_CAN_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # PDF Export
    if st.button("📄 Download Report as PDF"):
        buffer_pdf = BytesIO()
        c = canvas.Canvas(buffer_pdf, pagesize=A4)
        width, height = A4

        c.setFont("Helvetica-Bold", 16)
        c.drawString(40, height - 50, "EGYPTAIR M&E")
        c.setFont("Helvetica", 12)
        c.drawString(40, height - 70, "Technical Services Directorate")
        c.drawString(40, height - 90, "Reliability Department")
        c.drawString(40, height - 110, f"Date: {datetime.today().strftime('%Y-%m-%d')}")

        c.setFont("Helvetica-Bold", 14)
        c.drawString(40, height - 150, "CAN Dashboard Report")

        c.setFont("Helvetica", 11)
        y = height - 180
        c.drawString(40, y, f"Filters Used:")
        y -= 20
        c.drawString(60, y, f"• Years: {', '.join(map(str, selected_years))}")
        y -= 20
        c.drawString(60, y, f"• Quarters: {', '.join(map(str, selected_quarters))}")
        y -= 20
        c.drawString(60, y, f"• A/C Types: {', '.join(selected_types)}")

        y -= 40
        c.drawString(40, y, f"Summary:")
        y -= 20
        c.drawString(60, y, f"• Total Records: {len(filtered_df)}")
        y -= 20
        c.drawString(60, y, f"• Unique ATA Chapters: {filtered_df['ATA'].nunique()}")

        c.showPage()
        c.save()
        buffer_pdf.seek(0)

        st.download_button(
            label="📥 Click to download report (.pdf)",
            data=buffer_pdf,
            file_name="EGYPTAIR_CAN_Report.pdf",
            mime="application/pdf"
        )
