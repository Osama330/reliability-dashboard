import streamlit as st
import pandas as pd
import plotly.express as px
from datetime import datetime
from io import BytesIO
from docx import Document

# Ø¥Ø¹Ø¯Ø§Ø¯ Ø§Ù„ØµÙØ­Ø©
st.set_page_config(page_title="CAN Dashboard", layout="wide")

# CSS Ù„Ù„Ø®Ù„ÙÙŠØ©
st.markdown("""
    <style>
    body {
        background-color: #f2f7fa;
    }
    .block-container {
        padding-top: 20px;
    }
    </style>
""", unsafe_allow_html=True)

# Ø±Ø£Ø³ Ø§Ù„ØµÙØ­Ø© Ø§Ù„Ø±Ø³Ù…ÙŠ
col1, col2 = st.columns([3, 1])
with col1:
    st.markdown(f"""
        <div style='text-align: left; font-size: 22px; font-weight: bold; line-height: 1.6;'>
            EGYPTAIR M&E<br>
            Technical Services Directorate<br>
            Reliability Department<br>
            Date: <span style='font-weight: normal;'>{datetime.today().strftime('%Y-%m-%d')}</span>
        </div>
    """, unsafe_allow_html=True)
with col2:
    st.image("egyptair_logo.png", width=180)

# Ø¹Ù†ÙˆØ§Ù† Ø§Ù„Ø¯Ø§Ø´Ø¨ÙˆØ±Ø¯
st.title("ğŸ› ï¸ Component Alert Notice (CAN) Dashboard")

# ØªØ­Ù…ÙŠÙ„ ÙˆØªÙ†Ø¸ÙŠÙ Ø§Ù„Ø¨ÙŠØ§Ù†Ø§Øª
df_raw = pd.read_excel("CAN.xlsx")
df = df_raw.loc[:, ~df_raw.columns.str.contains("^Unnamed")]
df.columns = df.columns.str.strip()

# Ø§Ù„ÙÙ„Ø§ØªØ±
years = sorted(df["YEAR"].dropna().unique())
quarters = sorted(df["QUARTER NO"].dropna().unique())
ac_types = sorted(df["A/C TYPE"].dropna().unique())

left_col, right_col = st.columns([1, 2])
with left_col:
    st.subheader("ğŸ” Filters")
    selected_years = st.multiselect("Select Year(s)", years, default=years)
    selected_quarters = st.multiselect("Select Quarter(s)", quarters, default=quarters)
    selected_types = st.multiselect("Select A/C Type(s)", ac_types, default=ac_types)

# ÙÙ„ØªØ±Ø© Ø§Ù„Ø¨ÙŠØ§Ù†Ø§Øª
filtered_df = df[
    df["YEAR"].isin(selected_years) &
    df["QUARTER NO"].isin(selected_quarters) &
    df["A/C TYPE"].isin(selected_types)
]

# Ø¹Ø±Ø¶ Ø§Ù„Ø¨ÙŠØ§Ù†Ø§Øª
with right_col:
    st.markdown("### Filtered CAN Data")
    st.dataframe(filtered_df, use_container_width=True)

    # Ø²Ø± ØªØ­Ù…ÙŠÙ„ Excel
    excel_buffer = BytesIO()
    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
        filtered_df.to_excel(writer, sheet_name="Filtered CAN", index=False)
    excel_buffer.seek(0)
    st.download_button(
        label="ğŸ“¥ Download Filtered Data as Excel",
        data=excel_buffer,
        file_name="Filtered_CAN_Data.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

    # ===== Chart 1 =====
    st.subheader("ğŸ”§ Number of Removals by ATA")
    sort_removal = st.selectbox("Sort Order (Removals)", ["Descending", "Ascending"], index=0)
    top_removal = st.selectbox("Show Top (Removals)", ["All", "Top 3", "Top 6", "Top 10"], index=0)

    removal_by_ata = filtered_df.groupby("ATA")["NO OF REMOVAL"].sum().reset_index()
    removal_by_ata = removal_by_ata.sort_values("NO OF REMOVAL", ascending=(sort_removal == "Ascending"))
    if top_removal != "All":
        n = int(top_removal.split()[1])
        removal_by_ata = removal_by_ata.head(n)

    fig1 = px.bar(
        removal_by_ata,
        x="ATA",
        y="NO OF REMOVAL",
        text="NO OF REMOVAL",
        color_discrete_sequence=["#1f77b4"],
        category_orders={"ATA": removal_by_ata["ATA"].tolist()}
    )
    fig1.update_traces(textposition="outside")
    fig1.update_layout(
        yaxis_tickformat=",d",
        xaxis_tickmode='linear',
        xaxis_tickangle=-45
    )
    st.plotly_chart(fig1, use_container_width=True)

    # ===== Chart 2 =====
    st.subheader("ğŸ“Š Average Removal Rate by ATA")
    sort_rate = st.selectbox("Sort Order (Rate)", ["Descending", "Ascending"], index=0)
    top_rate = st.selectbox("Show Top (Rates)", ["All", "Top 3", "Top 6", "Top 10"], index=0)

    removal_rate_by_ata = filtered_df.groupby("ATA")["REMOVAL RATE"].mean().reset_index()
    removal_rate_by_ata = removal_rate_by_ata.sort_values("REMOVAL RATE", ascending=(sort_rate == "Ascending"))
    if top_rate != "All":
        n = int(top_rate.split()[1])
        removal_rate_by_ata = removal_rate_by_ata.head(n)

    fig2 = px.bar(
        removal_rate_by_ata,
        x="ATA",
        y="REMOVAL RATE",
        text="REMOVAL RATE",
        color_discrete_sequence=["#ff7f0e"],
        category_orders={"ATA": removal_rate_by_ata["ATA"].tolist()}
    )
    fig2.update_traces(texttemplate='%{text:.2f}', textposition="outside")
    fig2.update_layout(
        xaxis_tickmode='linear',
        xaxis_tickangle=-45
    )
    st.plotly_chart(fig2, use_container_width=True)

    # ===== Chart 3 =====
    st.subheader("ğŸ“ˆ CAN Distribution by ATA (Pie Chart)")
    top_pie = st.selectbox("Show Top (Pie)", ["All", "Top 3", "Top 6", "Top 10"], index=0)

    removal_by_ata_pie = filtered_df.groupby("ATA")["NO OF REMOVAL"].sum().reset_index()
    removal_by_ata_pie = removal_by_ata_pie.sort_values("NO OF REMOVAL", ascending=False)
    if top_pie != "All":
        n = int(top_pie.split()[1])
        removal_by_ata_pie = removal_by_ata_pie.head(n)

    if not removal_by_ata_pie.empty:
        fig3 = px.pie(
            removal_by_ata_pie,
            names="ATA",
            values="NO OF REMOVAL",
            color_discrete_sequence=px.colors.sequential.Blues
        )
        fig3.update_traces(textinfo='label+percent+value')
        st.plotly_chart(fig3, use_container_width=True)
    else:
        st.warning("No data available to display the Pie Chart.")

    # ===== Chart 4 =====
    st.subheader("ğŸ“Œ Number of CAN Entries per ATA Chapter")
    sort_can = st.selectbox("Sort Order (CAN Count)", ["Descending", "Ascending"], index=0)
    top_can = st.selectbox("Show Top (CAN Count)", ["All", "Top 3", "Top 6", "Top 10"], index=0)

    can_count_by_ata = filtered_df.groupby("ATA").size().reset_index(name="CAN Count")
    can_count_by_ata = can_count_by_ata.sort_values("CAN Count", ascending=(sort_can == "Ascending"))
    if top_can != "All":
        n = int(top_can.split()[1])
        can_count_by_ata = can_count_by_ata.head(n)

    if not can_count_by_ata.empty:
        fig4 = px.bar(
            can_count_by_ata,
            x="CAN Count",
            y="ATA",
            orientation="h",
            text="CAN Count",
            color_discrete_sequence=["#2ca02c"],
            category_orders={"ATA": can_count_by_ata["ATA"].tolist()}
        )
        fig4.update_traces(textposition="outside")
        fig4.update_layout(
            xaxis_tickformat=",d",
            yaxis_tickmode='linear',
            yaxis_tickfont=dict(size=12)
        )
        st.plotly_chart(fig4, use_container_width=True)
    else:
        st.warning("No data available to display CAN count chart.")

# ===== Export Word Report =====
st.markdown("---")
st.subheader("ğŸ“¥ Export Report")

if st.button("ğŸ“„ Download Word Report"):
    doc = Document()
    doc.add_heading("EGYPTAIR M&E", level=1)
    doc.add_paragraph("Technical Services Directorate")
    doc.add_paragraph("Reliability Department")
    doc.add_paragraph(f"Date: {datetime.today().strftime('%Y-%m-%d')}")
    doc.add_paragraph("Component Alert Notice (CAN) Report")

    doc.add_paragraph(f"Filters Used:")
    doc.add_paragraph(f"  â€¢ Years: {', '.join(map(str, selected_years))}")
    doc.add_paragraph(f"  â€¢ Quarters: {', '.join(map(str, selected_quarters))}")
    doc.add_paragraph(f"  â€¢ A/C Types: {', '.join(selected_types)}")

    doc.add_paragraph("\nSummary Data:")
    doc.add_paragraph(f"  â€¢ Total Records: {len(filtered_df)}")
    doc.add_paragraph(f"  â€¢ Unique ATA Chapters: {filtered_df['ATA'].nunique()}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="ğŸ“¥ Click to download report (.docx)",
        data=buffer,
        file_name="EGYPTAIR_CAN_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
